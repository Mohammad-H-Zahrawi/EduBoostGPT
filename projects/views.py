import openai
import os

os.environ['OPENAI_API_KEY']= 'sk-gdyx6VplClkZSi'
os.getenv('OPENAI_API_KEY')
openai.api_key = os.getenv('OPENAI_API_KEY')

from django.shortcuts import render, redirect 
from django.urls import reverse
from projects.models import TrackTeacherModel, ExamGPTModel, LessonPlanGPTModel
from projects.models import ExamGPTWhisperModel, LessonGPTWhisperModel
from django.contrib import messages
import pandas as pd
import numpy as np
pd.options.mode.chained_assignment = None  # default='warn'
import docxtpl
from docxtpl import DocxTemplate
from django.http import HttpResponse
from docx import Document

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders


import docxtpl
from docxtpl import DocxTemplate
from datetime import date, datetime

global exam_path
global exam_path_arabic

global answers_path

#Using Bard APIs
# from bardapi import Bard
# #Replace XXXX with the values you get from __Secure-1PSID
# os.environ['_BARD_API_KEY']="ZAjh4qajOjkdj2hDildkr_2i3mMPKY-eYx3C6BTfGLDv6SY06YYFmPmhG87J1oFaGG_nPw."

# import requests
# from bardapi.constants import SESSION_HEADERS
# from bardapi import Bard

# token = "aQjh4oLDS577NwiuWce_NnZSnBIzFDrZAX81tlUNFn_OxKaZFrJc-xQJwCwKM448Td9W0g."
# session = requests.Session()
# session.headers = SESSION_HEADERS
# session.cookies.set("__Secure-1PSID", token)
# session.cookies.set("__Secure-1PSIDTS", "sidts-CjIBSAxbGXZloT1-ncoPFGy5qMPzAYczSjyGWoydXdnAip17kQU8sh6TlRQ2JI8_OMGGJRAA")
# session.cookies.set("__Secure-1PSIDCC", "APoG2W8wQn_T3mgEthQdz0XaGwBztbmh_OGl9UE5XtpiSxIYMDOh75nka-bE5ga6gEZ79RobeQ")
# bard = Bard(token=token, session=session)



def track_teacher(request):
    if request.method == "POST":
        old=(request.FILES['oldfile'])
        new=(request.FILES['newfile'])
        TrackTeacherModel.objects.create(oldfile=old, newfile=new).save()
        messages.success(request, "File uploaded successfully!")
            
        old_original = pd.read_excel(old)
        new_original = pd.read_excel(new)
        
        new_original  = new_original.sort_values('Email')
        old_original = old_original.set_index('Email')
        old_original = old_original[~old_original.index.duplicated()]
        old_original = old_original.reindex(index=new_original['Email'])
        old_original = old_original.reset_index()
        new_original = new_original.reset_index()

        columns = ['School Name', 'Full Name', 'Created Content', 'Created Discussions', 'No Of Posts', 'Question Counts In QB', 'Questions In Assessment', 'Assessments Count', 'Quizzes Count', 'InClass Activity Count', 'Added Materials On Space']
        old = old_original[columns]
        new = new_original[columns]
        old['total'] = old.iloc[:,2:].sum(axis=1)
        new['total'] = new.iloc[:,2:].sum(axis=1)
        columns.append('total')
        df_difference = pd.DataFrame()
        df_difference['School Name'] = new['School Name']
        df_difference['Full Name'] = new['Full Name']
        df_difference = df_difference.dropna()

        for i in columns[2:]:
            df_difference[i] = new[i] - old[i]

        # new = new.reset_index(drop=True)
        # new['Full Name'] = new['Full Name'].apply(lambda x:GoogleTranslator(source='en', target='ar').translate(x))

        for i, x in enumerate(new['Full Name']):
            if len(x.split(' '))>2:
                list_names = x.split(' ')
                new['Full Name'][i] = ' '.join([list_names[0], list_names[1],  list_names[-1]])
                
        df_difference['Full Name'] = new['Full Name']
        df_difference['total_new'] = new['total']
        df_difference['total_old'] = old['total']
        df_difference['Total Points'] = new_original['Total Points']
        
        df_sorted_total=df_difference.sort_values('Total Points')
        df_sorted_total_names  =df_sorted_total['Full Name']
        total_points = df_sorted_total['Total Points']
        total_new = df_sorted_total['total_new']
        total_old = df_sorted_total['total_old']
        difference = df_sorted_total['total']
        
        track_teachers_report = df_difference.copy()
        track_teachers_report = track_teachers_report[['Full Name', 'total']]
        track_teachers_report.columns = ['Full Name', 'Improvement Index' ]
        track_teachers_report['KG'] = track_teachers_report['Improvement Index'].apply(lambda x:'yes' if x>50 else 'No')
        track_teachers_report['C1'] = track_teachers_report['Improvement Index'].apply(lambda x:'yes' if x>100 else 'No')
        track_teachers_report['C2/C3'] = track_teachers_report['Improvement Index'].apply(lambda x:'yes' if x>150 else 'No')
        
        global track_teachers_report_path
        track_teachers_report_name = 'Report teacher performance' + ' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')
        track_teachers_report_path = 'media/track_teachers/'+track_teachers_report_name+'.xlsx'
        track_teachers_report.to_excel(track_teachers_report_path)
        if os.path.exists(track_teachers_report_path):
            report_exist = True

        background=[]
        background_difference = ['rgba(50, 255, 50, 1' if x>0 else 'rgba(255, 100, 100, 1' for x in df_sorted_total['total'] ]

        number_rows = len(df_difference)
        for x in range(number_rows):
            if x < np.round(number_rows/5):
                value = 'rgba(255, 0, 0, 1'
            elif x <np.round(number_rows/5)*2:
                value = 'rgba(255, 100, 100, 1'
            elif x <np.round(number_rows/5)*3:
                value = 'rgba(255, 255, 50, 1'
            elif x <np.round(number_rows/5)*4:
                value = 'rgba(50, 255, 50, 1'
            else:
                value = 'rgba(0, 150, 0, 1'

            background.append(value)
        
        global email_list
        email_list='engmohammadzahrawi1996@gmail.com,'
        for i in list(new_original['Email']):
            email_list = email_list+i+','
        new['% Created Content'] = (new['Created Content'] / new['total']) *100
        new['% Discussions'] = (new['Created Discussions'] / new['total']) *100
        new['% Posts'] = (new['No Of Posts'] / new['total']) *100

        new['% Questions in Banks'] = (new['Question Counts In QB'] / new['total']) *100
        new['% Questions in Assessments'] = (new['Questions In Assessment'] / new['total']) *100
        new['% assessments'] = (new['Assessments Count'] / new['total']) *100
        new['% quizzes'] = (new['Quizzes Count'] / new['total']) *100

        new['% Activity'] = (new['InClass Activity Count'] / new['total']) *100
        new['% Adding Materials'] = (new['Added Materials On Space'] / new['total']) *100

        pieChart_data = new[['% Created Content', '% Discussions', '% Posts','% Questions in Banks', '% Questions in Assessments', '% assessments', '% quizzes', '% Activity', '% Adding Materials' ]]
        pieChart_names= new['Full Name']
        data=[]
        for i in range(len(new)):
            data.append([round(x, 2) for x in list(pieChart_data.loc[i])])
        data = np.array(data)
        data = np.nan_to_num(data,0)
        data = data.tolist()
            
        pieChart_IDs = ['ID_' + str(x) for x in range(len(new))]
        config_id = ['config_' + str(x) for x in range(len(new))]
        data_id = ['data_' + str(x) for x in range(len(new))]
        combine = zip(pieChart_IDs, config_id, data_id)
        
        
        context = {'total_points':total_points, 'df_sorted_total_names':df_sorted_total_names, 'total_new':total_new, 
                   'total_old':total_old, 'background': background, 'background_difference':background_difference, 
                    'difference':difference, 'data' : data, 'pieChart_IDs':pieChart_IDs, 'pieChart_names':pieChart_names, 'combine':combine, 
                    'report_exist':report_exist}

        # context = {'total_points':total_points, 'df_sorted_total_names':df_sorted_total_names}
        return render(request, 'projects/track_teacher.html', context =context)
    
    else:
        context = {}
        return render(request, 'projects/track_teacher.html', context = context)
    
    
# def thank(request): 
#     if request.method == "POST":
#         teacher = request.POST.get('teacher_name')
        
#         GDRAT_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),'projects/media/ex.docx')
#         filee = DocxTemplate(GDRAT_abs_path.replace('\\', '/'))
#         co = {'teacher':teacher}
#         filee.render(co)
#         filee.save('a.docx')
        
#         context = {'teacher':teacher}
#         return render(request, 'projects/thank.html', context=context)
#     else:
#         return render(request, 'projects/thank.html')
        
# import mimetypes
# from django.http.response import HttpResponse


# def draft1(request):
#     base = os.path.dirname(os.path.dirname(os.path.abspath(__file__))).replace('\\', '/')
#     base = base + '/a.docx'
#     content_type_value = 'text/plain'
#     with open(base, 'rb') as fh:
#         response = HttpResponse(
#                 fh.read(),
#                 content_type=content_type_value
#             )

#     response['Content-Disposition'] = "attachment; filename=" + base
#     return response


def examGPT(request):




    if request.method  == 'POST':
        school_name = request.POST.get('school_name')
        grade = request.POST.get('grade')
        subject = request.POST.get('subject')
        lesson_name = request.POST.get('lesson_name')
        teacher_name = request.POST.get('teacher_name')
        topic_language = request.POST.get('topic_language')
        ExamGPTModel.objects.create(school_name=school_name, grade=grade, 
                                    subject=subject,lesson_name=lesson_name,teacher_name=teacher_name, topic_language=topic_language).save()





        if topic_language=='English':

            global exam_path
            global answers_path
            exam_path = '#'
            answers_path = '#'
        
            exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),'projects/media/examGPT/exam_template.docx')
            exam_abs_path = exam_abs_path.replace('\\', '/')
            exam = DocxTemplate(exam_abs_path)
            today = date.today()

            # using chatGPT API

            
            def create_test_prompt(topic):
                prompt = f"create a multiple choice quiz on the topic of {topic} consisting of five questions with correct answers."
                return prompt

            response = openai.Completion.create(
                model = 'text-davinci-003', 
                prompt = create_test_prompt(lesson_name),  # one vairables
                max_tokens = 2500, 
                temperature = 0.7)
            
            output_text = (response['choices'][0]['text'])

            # output_text = "\n\n1. What is the result of subtracting -7 from 9?\nA. 16\nB. 2\nC. -2\nD. -16\n\nAnswer: B. 2\n\n2. What is the result of adding -3 and -5?\nA. -8\nB. 8\nC. -2\nD. 2\n\nAnswer: A. -8\n\n3. What is the result of subtracting -12 from -5?\nA. -7\nB. 7\nC. -17\nD. 17\n\nAnswer: D. 17\n\n4. What is the result of adding 5 and -2?\nA. 3\nB. -7\nC. -3\nD. 7\n\nAnswer: A. 3\n\n5. What is the result of subtracting -4 from -6?\nA. -2\nB. 10\nC. 2\nD. -10\n\nAnswer: C. 2"
            
            def create_student_view(test):
                student_view = {1:''}
                question_number = 1
                for line in test.split('\n'):
                    if not line.startswith('Answer:') and not line.startswith('Correct Answer'):
                        student_view[question_number] += line + '\n'
                    else:
                        question_number +=1
                        student_view[question_number] = ''
                return student_view


            def extract_answers(test):
                answers = {1:''}
                question_number = 1
                for line in test.split('\n'):
                    if line.startswith('Answer:') or line.startswith('Correct Answer'):
                        answers[question_number] += line + '\n'
                        question_number +=1
                        answers[question_number] = ''
                return answers
            
            questions = create_student_view(output_text)
            answers = extract_answers(output_text)
            variables_needed_to_changed = {'Q1':questions[1][2:-1], 'Q2':questions[2][1:-1], 'Q3':questions[3][1:-1], 'Q4':questions[4][1:-1], 'Q5': questions[5][1:-1], 'subject':subject, 'SchoolName':school_name, 'grade':grade, 'date':str(today), 'teacher':teacher_name}
            exam.render(variables_needed_to_changed)
            answer_file = Document()
            for i in answers:
                answer_file.add_paragraph(f'Q{i}.{answers[i]}')
                if i==5:
                    break

            
            filename = 'Exam about ' + lesson_name + ' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')
            answers_path = 'Answers for ' + lesson_name + ' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')

            exam_path = 'media/examGPT/'+filename+'.docx'
            answers_path = 'media/examGPT/'+answers_path+'.docx'

            exam.save(exam_path)
            answer_file.save(answers_path)
            
            if os.path.exists(exam_path):
                exam_ready_for_downloading = True
            if os.path.exists(answers_path):
                answers_ready_for_downloading = True
            context={'exam_ready_for_downloading':exam_ready_for_downloading , 'answers_ready_for_downloading':answers_ready_for_downloading}

            return render(request, 'projects/ExamGPT.html', context=context)
        
        else: #if language arabic
            global exam_path_arabic
            exam_path_arabic = '#'
    
            exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),'projects/media/examGPT/exam_template_arabic.docx')
            exam_abs_path = exam_abs_path.replace('\\', '/')
            exam = DocxTemplate(exam_abs_path)
            today = date.today()

            ## using chatGPT

            
            def create_test_prompt(topic):
                prompt = f"قم بإنشاء اختبار متعدد الاختيارات حول موضوع  { topic } يتكون من خمسة أسئلة بإجابات صحيحة."
                return prompt

            response = openai.Completion.create(
                model = 'text-davinci-003', 
                prompt = create_test_prompt(lesson_name),  # two vairables
                max_tokens = 2500, 
                temperature = 0.7)
            
            output_text = (response['choices'][0]['text'])






            def create_student_view(text):
                student_view = {'questions':''}
                for line in text.split('\n')[2:-2]:
                    student_view['questions'] += line + '\n'
                return student_view

            questions = create_student_view(output_text)
            variables_needed_to_changed = {'Q1':questions['questions'], 'subject':subject, 'SchoolName':school_name, 'grade':grade, 'date':str(today), 'teacher':teacher_name}
            exam.render(variables_needed_to_changed)
            filename = 'Exam written in arabic'  +' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')
            exam_path_arabic = 'media/examGPT/'+filename+'.docx'
            exam.save(exam_path_arabic)
            if os.path.exists(exam_path_arabic):
                exam_ready_for_downloading_ar = True
            context={'exam_ready_for_downloading_ar':exam_ready_for_downloading_ar }

            return render(request, 'projects/ExamGPT.html', context=context)
        
    else:
        return render(request, 'projects/ExamGPT.html')

from django.http import JsonResponse

# import sounddevice as sd
# from scipy.io.wavfile import write
# import wavio as wv


# def record(request):
#     import pyaudio
#     import wave

#     # define the audio settings
#     chunk = 1024  # number of audio samples per buffer
#     sample_format = pyaudio.paInt16  # audio sample format
#     channels = 1  # mono audio
#     sample_rate = 44100  # number of audio samples per second
#     duration = 6  # duration of recording in seconds

#     # create the audio object
#     audio = pyaudio.PyAudio()

#     # open the microphone stream
#     stream = audio.open(format=sample_format, channels=channels,
#                         rate=sample_rate, frames_per_buffer=chunk,
#                         input=True)

#     # start recording
#     frames = []
#     for i in range(int(sample_rate / chunk * duration)):
#         data = stream.read(chunk)
#         frames.append(data)

#     # stop recording
#     stream.stop_stream()
#     stream.close()
#     audio.terminate()

#     # save the audio to a WAV file
#     global gpt_voice_record_path
#     gpt_voice_record_name = 'GPT record ' + datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')
#     gpt_voice_record_path = 'media/examGPT_VoiceRecordings/'+gpt_voice_record_name+'.wav'
            
#     wave_file = wave.open(gpt_voice_record_path, 'wb')
#     wave_file.setnchannels(channels)
#     wave_file.setsampwidth(audio.get_sample_size(sample_format))
#     wave_file.setframerate(sample_rate)
#     wave_file.writeframes(b''.join(frames))
#     wave_file.close()
#     voice_recorded_completed  = True
#     context = {'voice_recorded_completed':voice_recorded_completed}
#     return render(request, 'projects/ExamGPT_Voice.html', context=context)




    
# def examGPT_voice(request):
#     if request.method=='POST':
#         topic_language = request.POST.get('topic_language')
#         global exam_path_whisper
#         global answers_path_whisper
        
        
#         if topic_language=='English':
#             openai.api_key = "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"
#             audio_file = open(gpt_voice_record_path, "rb")
#             transcript = openai.Audio.transcribe("whisper-1", audio_file)
#             topic_name = transcript.text
#             ExamGPTWhisperModel.objects.create(lesson_name = topic_name ).save()

#             exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),'projects/media/examGPT/bard/exam_template.docx')
#             exam_abs_path = exam_abs_path.replace('\\', '/')
#             exam = DocxTemplate(exam_abs_path)
#             today = date.today()
            
#             def create_test_prompt(topic):
#                 prompt = f"create a multiple choice quiz on the topic of {topic} consisting of ten questions with correct answers, separating questions by newlines"
#                 return prompt

            # response = openai.Completion.create(
            #     model = 'text-davinci-003', 
            #     prompt = create_test_prompt(topic_name),  # one vairables
            #     max_tokens = 500, 
            #     temperature = 0.7)
            
            # output_text = (response['choices'][0]['text'])
            
            # def create_student_view(test):
            #     student_view = {1:''}
            #     question_number = 1
            #     for line in test.split('\n'):
            #         if not line.startswith('Answer:') and not line.startswith('Correct Answer'):
            #             student_view[question_number] += line + '\n'
            #         else:
            #             question_number +=1
            #             student_view[question_number] = ''
            #     return student_view


            # def extract_answers(test):
            #     answers = {1:''}
            #     question_number = 1
            #     for line in test.split('\n'):
            #         if line.startswith('Answer:') or line.startswith('Correct Answer'):
            #             answers[question_number] += line + '\n'
            #             question_number +=1
            #             answers[question_number] = ''
            #     return answers
            
            # questions = create_student_view(output_text)
            # answers = extract_answers(output_text)
            # variables_needed_to_changed = {'Q1':questions[1][2:-1], 'Q2':questions[2][1:-1], 'Q3':questions[3][1:-1], 'Q4':questions[4][1:-1], 'Q5': questions[5][1:-1], 'date':str(today)}
            # exam.render(variables_needed_to_changed)
            # answer_file = Document()
            # for i in answers:
            #     answer_file.add_paragraph(f'Q{i}.{answers[i]}')
            #     if i==5:
            #         break
            


            # filename = 'Exam about' + topic_name + ' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')
            # answers_path_whisper_name = 'Answers for ' + topic_name + ' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')

            # exam_path_whisper = 'media/examGPT_whisper/'+filename+'.docx'
            # answers_path_whisper = 'media/examGPT_whisper/'+answers_path_whisper_name+'.docx'

            # exam.save(exam_path_whisper)
            # answer_file.save(answers_path_whisper)
            
            
            # if os.path.exists(exam_path_whisper):
            #     exam_ready_for_downloading = True
            # if os.path.exists(answers_path_whisper):
            #     answers_ready_for_downloading = True
            # context={'exam_ready_for_downloading':exam_ready_for_downloading , 'answers_ready_for_downloading':answers_ready_for_downloading}
            
        #     output_text = bard.get_answer(create_test_prompt(topic_name))['content']

        #     import re
        #     def extract_answers(text):
        #         answers = {1:''}
        #         question_number = 1
        #         for line in text.split('\n')[2:-2]:
        #             if 'Answer' in line or 'correct answer' in line.lower():
        #                 answers[question_number] += ' '.join(re.findall('\w*', line)) + '\n'
        #                 question_number +=1
        #                 answers[question_number] = ''
        #         return answers

        #     def create_student_view(text):
        #         student_view = {'questions':''}
        #         for line in text.split('\n')[2:-2]:
        #             if not 'Answer' in line and not 'correct answer' in line.lower():
        #                 student_view['questions'] += line + '\n'
        #             else:
        #                 pass
        #         return student_view
            

        #     questions = create_student_view(output_text)
        #     answers  = extract_answers(output_text)

        #     answer_file = Document()
        #     for i in answers:
        #         answer_file.add_paragraph(f'Q{i}.{answers[i]}')
        #         if i==5:
        #             break
            
        #     variables_needed_to_changed = {'Q1':questions['questions']}
        #     exam.render(variables_needed_to_changed)

        #     filename = 'Exam about ' + topic_name + ' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')
        #     answers_path_whisper_name = 'Answers for ' + topic_name + ' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')

        #     exam_path_whisper = 'media/examGPT_whisper/'+filename+'.docx'
        #     answers_path_whisper = 'media/examGPT_whisper/'+answers_path_whisper_name+'.docx'

        #     exam.save(exam_path_whisper)
        #     answer_file.save(answers_path_whisper)
            
        #     if os.path.exists(exam_path_whisper):
        #         exam_ready_for_downloading = True
        #     if os.path.exists(answers_path_whisper):
        #         answers_ready_for_downloading = True
        #     context={'exam_ready_for_downloading':exam_ready_for_downloading , 'answers_ready_for_downloading':answers_ready_for_downloading}


        #     return render(request, 'projects/ExamGPT_Voice.html', context=context)
        
        # else:
        #     openai.api_key = "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"

        #     audio_file = open(gpt_voice_record_path, "rb")
        #     transcript = openai.Audio.transcribe("whisper-1", audio_file)
        #     topic_name = transcript.text
        #     ExamGPTWhisperModel.objects.create(lesson_name = topic_name ).save()

        #     exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),'projects/media/examGPT/bard/exam_template_arabic.docx')
        #     exam_abs_path = exam_abs_path.replace('\\', '/')
        #     exam = DocxTemplate(exam_abs_path)
        #     today = date.today()
            
        #     def create_test_prompt(topic):
        #         prompt = f"قم بإنشاء اختبار متعدد الاختيارات حول موضوع  { topic } يتكون من عشرة أسئلة بإجابات صحيحة."
        #         return prompt

            # response = openai.Completion.create(
            #     model = 'text-davinci-003', 
            #     prompt = create_test_prompt(topic_name),  # two vairables
            #     max_tokens = 800, 
            #     temperature = 0.7)
            
            # output_text = (response['choices'][0]['text'])

#             output_text = bard.get_answer(create_test_prompt(topic_name))['content']

#             def create_student_view(text):
#                 student_view = {'questions':''}
#                 for line in text.split('\n')[2:-2]:
#                     student_view['questions'] += line + '\n'
#                 return student_view

#             questions = create_student_view(output_text)
#             variables_needed_to_changed = {'Q1':questions['questions'], 'date':str(today)}
#             exam.render(variables_needed_to_changed)

    
#             filename = 'Exam written in arabic'  +' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')
            
#             exam_path_whisper = 'media/examGPT_whisper/'+filename+'.docx'

#             exam.save(exam_path_whisper)
#             if os.path.exists(exam_path_whisper):
#                 exam_ready_for_downloading = True
#             context={'exam_ready_for_downloading':exam_ready_for_downloading }

#             return render(request, 'projects/ExamGPT_Voice.html', context=context)
            
            
#     return render(request, 'projects/ExamGPT_Voice.html')
        
# def download_exam_whisper(request):
#         exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),exam_path_whisper)
#         exam_abs_path = exam_abs_path.replace('\\', '/')
#         if os.path.exists(exam_abs_path):
#             with open(exam_abs_path, 'rb') as fh:
#                 response = HttpResponse(fh.read(), content_type='application/octet-stream')
#                 response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(exam_path_whisper)
#                 return (response)       
            
# def download_exam_answers_whisper(request):
#         exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),answers_path_whisper)
#         exam_abs_path = exam_abs_path.replace('\\', '/')
#         if os.path.exists(exam_abs_path):
#             with open(exam_abs_path, 'rb') as fh:
#                 response = HttpResponse(fh.read(), content_type='application/octet-stream')
#                 response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(answers_path_whisper)
#                 return (response)  
        
def download_exam(request):
        exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),exam_path)
        exam_abs_path = exam_abs_path.replace('\\', '/')
        if os.path.exists(exam_abs_path):
            with open(exam_abs_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type='application/octet-stream')
                response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(exam_abs_path)
    
                return (response)
            
            
            
def download_exam_arabic(request):
        exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),exam_path_arabic)
        exam_abs_path = exam_abs_path.replace('\\', '/')
        if os.path.exists(exam_abs_path):
            with open(exam_abs_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type='application/octet-stream')
                response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(exam_abs_path)
    
                return (response)

            

def download_exam_answers(request):
        answers_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),answers_path)
        answers_abs_path = answers_abs_path.replace('\\', '/')
        with open(answers_abs_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type='application/octet-stream')
            response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(answers_abs_path)

            return (response)
        
        
def download_teachers_report(request):
        track_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),track_teachers_report_path)
        track_path = track_path.replace('\\', '/')
        with open(track_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type='application/octet-stream')
            response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(track_path)

            return (response)
        
        
        
        
def send_trackTeachers(request):
    answers_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),track_teachers_report_path)
    answers_abs_path = answers_abs_path.replace('\\', '/')  

    def gamil_login(username, password):
        smtp_server_name = 'smtp.gmail.com'
        port_number = 587
        server = smtplib.SMTP(host=smtp_server_name, port = port_number)
        server.connect(host=smtp_server_name, port=port_number)
        server.starttls()
        server.login(username, password=password)
        return server
    
    msg = MIMEMultipart()
    msg['subject']="auto-generated message"
    msg['From'] = 'engmohammadzahrawi1996@gmail.com'
    msg['To']= email_list
    msg['Bcc'] = ''    
    msg.set_charset('utf8')
    #email_list
    
    body = """
   <strong> Dear Teachers, <br>Please find the latest LMS report.<br> ALL Regards <br> EduBoostAI team <br> </strong>
    """
    msg.attach(MIMEText(body, 'html'))
    file_path = answers_abs_path
    file_name = 'LMS report_UAE.xlsx'
    part = MIMEBase('application', "octet-stream")
    part.set_payload(open(file_path, "rb").read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment', filename=file_name)
    msg.attach(part)
  
    # file = open(file_path, 'rb')
    # payload = MIMEText('text', 'csv')
    # payload.set_payload(file.read())
    # file.close()  
    # encoders.encode_base64(payload)
    # payload.add_header('Content-Disposition', 'attachment', filename= file_name)
    # msg.attach(payload)

    server = gamil_login(username='engmohammadzahrawi1996@gmail.com', password='ktkruwavoxxiboht')
    server.send_message(msg)
    server.quit()
    
    return render(request, 'projects/track_teacher_send_reports.html')


def signup(request):
    return render(request, 'projects/signup.html', context = {})

    
    
def LessonPlanGPT (request):
    if request.method == "POST":
        school_name = request.POST.get('school_name')
        grade = request.POST.get('grade')
        subject = request.POST.get('subject')
        lesson_name = request.POST.get('lesson_name')
        teacher_name = request.POST.get('teacher_name')
        LessonPlanGPTModel.objects.create(lesson_name = lesson_name, subject=subject, grade=grade, school_name=school_name, teacher_name=teacher_name ).save()
        lessonPlan_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),'projects/media/lessonPlanGPT/lessonPlanTemplate.docx')
        lessonPlan_abs_path = lessonPlan_abs_path.replace('\\', '/')
        lessonPlanFile = DocxTemplate(lessonPlan_abs_path)
        today = date.today()
        

        # # chatGPT
        # os.environ['OPENAI_API_KEY']= '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'
        # os.getenv('OPENAI_API_KEY')
        # openai.api_key = os.getenv('OPENAI_API_KEY')
        
        def create_lesson_prompt(topic):
            prompt = f"create a lesson plan about {topic} showing Objectives, Materials, Introduction, Direct Instruction, Guided Practice , Independent Practice, Closure, and Assessment."
            return prompt

        response = openai.Completion.create(
            model = 'text-davinci-003', 
            prompt = create_lesson_prompt(lesson_name),  # one vairables
            max_tokens = 2500, 
            temperature = 0.7)
        
        output_text = (response['choices'][0]['text'])
        # output_text="\n\nObjectives:\n  Students will be able to identify the consequences of protein deficiency in the body.\n\nMaterials:\n  -Whiteboard and markers\n  -PowerPoint presentation\n  -Handouts\n\nIntroduction (5 minutes):\n  -Introduce the topic of protein deficiency and explain why it is important to understand the consequences of a protein deficiency.\n\nDirect Instruction (15 minutes):\n  -Explain the functions of proteins in the body and how a protein deficiency can lead to health problems.\n  -Discuss the different types of proteins and how they affect the body.\n\nGuided Practice (15 minutes):\n  -Have the students break into small groups and discuss the consequences of a protein deficiency.\n  -Provide each group with a handout with a list of the consequences and have them discuss and identify which consequences they think are most severe.\n\nIndependent Practice (15 minutes):\n  -Have the students work independently to research the consequences of a protein deficiency and create a po"
        

        objective_flag, Materials_flag, Introduction_flag, Direct_flag = False, False, False, False
        Guided_flag, Independent_flag, Closure_flag, Assessment_flag = False, False, False, False
        objective, Materials, Introduction, Direct, Guided, Independent, Closure, Assessment = [], [], [], [], [], [], [], []

        #chatGPT code
        for line in output_text.split('\n'):
            if line =='':
                continue
            if line.startswith('Objectives'):
                objective_flag = True
                Materials_flag,Introduction_flag,Direct_flag,Guided_flag,Independent_flag,Closure_flag,Assessment_flag = False,False,False,False,False,False,False 
            elif line.startswith('Materials') and not line.startswith('\n'):
                Materials_flag = True
                objective_flag,Introduction_flag,Direct_flag,Guided_flag,Independent_flag,Closure_flag,Assessment_flag = False,False,False,False,False,False,False 
            elif line.startswith('Introduction'):
                Introduction_flag = True
                objective_flag,Materials_flag,Direct_flag,Guided_flag,Independent_flag,Closure_flag,Assessment_flag = False,False,False,False,False,False,False 
            elif line.startswith('Direct Instruction'):
                Direct_flag = True
                objective_flag,Materials_flag,Introduction_flag,Guided_flag,Independent_flag,Closure_flag,Assessment_flag = False,False,False,False,False,False,False 
            elif line.startswith('Guided Practice'):
                Guided_flag = True
                objective_flag, Materials_flag, Introduction_flag, Direct_flag, Independent_flag, Closure_flag, Assessment_flag = False,False,False,False,False,False,False
            elif line.startswith('Independent Practice'):
                Independent_flag = True
                objective_flag, Materials_flag, Introduction_flag, Direct_flag, Guided_flag, Closure_flag, Assessment_flag = False,False,False,False,False,False,False
            elif line.startswith('Closure'):
                Closure_flag = True
                objective_flag, Materials_flag, Introduction_flag, Direct_flag, Guided_flag, Independent_flag, Assessment_flag = False,False,False,False,False,False,False
            elif line.startswith('Assessment'):
                Assessment_flag = True
                objective_flag, Materials_flag, Introduction_flag, Direct_flag, Guided_flag, Independent_flag, Closure_flag = False,False,False,False,False,False,False
            else:
                pass
            if objective_flag:
                objective.append(line)      
            if Materials_flag:
                Materials.append(line)   
            if Introduction_flag:
                Introduction.append(line)      
            if Direct_flag:
                Direct.append(line) 
            if Guided_flag:
                Guided.append(line)      
            if Independent_flag:
                Independent.append(line)   
            if Closure_flag:
                Closure.append(line)      
            if Assessment_flag:
                Assessment.append(line)  




        objective = '\n'.join(objective[1:])
        Materials = '\n'.join(Materials[1:])
        Introduction = '\n'.join(Introduction)
        Direct = '\n'.join(Direct)
        Guided = '\n'.join(Guided)
        Independent = '\n'.join(Independent)
        Closure = '\n'.join(Closure)
        Assessment = '\n'.join(Assessment)
        variables_in_docFile = {'objective':objective, 'materials':Materials,
                                'introduction':Introduction, 'direct':Direct, 'guided': Guided, 
                                'independent':Independent, 'closure':Closure,
                                'assessment':Assessment , 'teacher':teacher_name, 'subject':subject, 'date':today, 
                                    'lesson':lesson_name, 'school_name':school_name, 'grade':grade}
        lessonPlanFile.render(variables_in_docFile)
        
        
        
        global lessonPlanFile_path
        lessonPlanFile_path='#'
        filename = 'Lesson Plan about ' + lesson_name + ' '+ datetime.strftime(datetime.now(), '%Y-%m-%d-%H-%M-%S')

        lessonPlanFile_path = 'media/lessonPlan_GPT/'+filename+'.docx'

        lessonPlanFile.save(lessonPlanFile_path)
        if os.path.exists(lessonPlanFile_path):
            lessonPlan_ready_for_downloading = True
        context={'lessonPlan_ready_for_downloading':lessonPlan_ready_for_downloading}
        
        return render(request, 'projects/LessonPlanGPT.html', context = context)   

    return render(request, 'projects/LessonPlanGPT.html')   


def download_lessonPlan(request):
        exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),lessonPlanFile_path)
        exam_abs_path = exam_abs_path.replace('\\', '/')
        if os.path.exists(exam_abs_path):
            with open(exam_abs_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type='application/octet-stream')
                response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(lessonPlanFile_path)
                return (response) 
            
# def download_lessonPlan(request):
#         exam_abs_path =  os.path.join(os.path.dirname(os.path.dirname(__file__)),lessonPlanFile_path_whisper)
#         exam_abs_path = exam_abs_path.replace('\\', '/')
#         if os.path.exists(exam_abs_path):
#             with open(exam_abs_path, 'rb') as fh:
#                 response = HttpResponse(fh.read(), content_type='application/octet-stream')
#                 response['Content-Disposition'] = 'attachment; filename=' + os.path.basename(exam_path_whisper)
#                 return (response)     
